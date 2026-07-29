import io

from docx import Document
from pypdf import PdfWriter

from app.integrations.drive.service import DriveService


def test_drive_search_escapes_quotes_and_backslashes() -> None:
    assert DriveService._search_query("Atlas's \\ plan") == (
        "(name contains 'Atlas\\'s \\\\ plan' or "
        "fullText contains 'Atlas\\'s \\\\ plan') and trashed = false"
    )


def test_drive_metadata_record_normalizes_provider_values() -> None:
    record = DriveService._metadata_record(
        {
            "id": "file-1",
            "name": "Atlas plan.pdf",
            "mimeType": "application/pdf",
            "size": "42",
            "modifiedTime": "2026-07-29T08:00:00.000Z",
            "parents": ["folder-1"],
        }
    )

    assert record["name"] == "Atlas plan.pdf"
    assert record["size"] == 42
    assert record["modified_time"].isoformat() == "2026-07-29T08:00:00+00:00"
    assert record["parents"] == ["folder-1"]


def test_drive_metadata_view_preserves_google_mime_key_for_content_dispatch() -> None:
    view = DriveService._metadata_view(
        {"id": "file-1", "name": "Atlas.pdf", "mimeType": "application/pdf"}
    )

    assert view["mimeType"] == "application/pdf"


def test_drive_reads_docx_text() -> None:
    stream = io.BytesIO()
    document = Document()
    document.add_heading("Atlas", level=1)
    document.add_paragraph("Drive integration works.")
    document.save(stream)

    assert DriveService._docx_text(stream.getvalue()) == (
        "Atlas\nDrive integration works."
    )


def test_drive_reads_pdf_without_text() -> None:
    stream = io.BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.write(stream)

    assert DriveService._pdf_text(stream.getvalue()) == ""
